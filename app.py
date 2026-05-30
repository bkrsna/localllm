"""
Flask Backend for Ollama LLM Chat Interface
Provides REST API endpoint to communicate with local Ollama instance
and stores chat history in SQLite database.
"""

from flask import Flask, request, jsonify, Response, stream_with_context, render_template, redirect, url_for, flash, send_file
from flask_cors import CORS
import requests
import json
import os
import io
import base64
import uuid
from datetime import datetime
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, login_user, login_required, logout_user, current_user

# Import models
from models import db, User, ChatSession, ChatMessage, Persona, Document, DocumentChunk

app = Flask(__name__)
CORS(app)

# Configuration
OLLAMA_API_URL = "http://localhost:11434/api/generate"
DEFAULT_MODEL = "gemma3:1b"
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(BASE_DIR, 'chat.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = 'dev-secret-key-change-in-prod' # simple key for dev

# Initialize Plugins
db.init_app(app)
login_manager = LoginManager()
login_manager.login_view = 'login'
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# --- Routes ---

@app.route('/')
def index():
    return render_template('home.html')

@app.route('/chat')
def chat_interface():
    # Allow guests
    return render_template('index.html')

@app.after_request
def add_header(response):
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '-1'
    return response

# === Auth Routes ===

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter_by(username=username).first()
        if user:
            flash('Username already exists')
            return redirect(url_for('signup'))
        
        new_user = User(
            username=username, 
            password_hash=generate_password_hash(password, method='scrypt')
        )
        # First user is admin automatically for simple bootstrap
        if User.query.count() == 0:
            new_user.is_admin = True
            
        db.session.add(new_user)
        db.session.commit()
        
        login_user(new_user)
        return redirect(url_for('chat_interface'))
        
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            return redirect(url_for('chat_interface'))
        else:
            flash('Invalid username or password')
            
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

@app.route('/admin')
@login_required
def admin():
    if not current_user.is_admin:
        return "Access Denied", 403
    
    users = User.query.all()
    user_data = []
    
    total_sessions = 0
    total_messages = 0
    
    for u in users:
        # Get sessions for this user
        u_sessions = ChatSession.query.filter_by(user_id=u.id).all()
        s_count = len(u_sessions)
        
        # Count messages
        m_count = 0
        for s in u_sessions:
            m_count += ChatMessage.query.filter_by(session_id=s.id).count()
            
        user_data.append({
            "username": u.username,
            "is_admin": u.is_admin,
            "session_count": s_count,
            "message_count": m_count
        })
        
        total_sessions += s_count
        total_messages += m_count

    # Also count guest sessions
    guest_sessions = ChatSession.query.filter_by(user_id=None).count()
    total_sessions += guest_sessions

    return render_template('admin.html', 
                         users=user_data, 
                         stats={
                             "total_sessions": total_sessions, 
                             "total_messages": total_messages,
                             "guest_sessions": guest_sessions
                         })

# === Persona Routes ===

@app.route('/api/personas', methods=['GET'])
def get_personas():
    system_p = Persona.query.filter_by(is_system=True).all()
    user_p = []
    if current_user.is_authenticated:
        user_p = Persona.query.filter_by(is_system=False, user_id=current_user.id).all()
    all_p = system_p + user_p
    return jsonify({"personas": [p.to_dict() for p in all_p]})

@app.route('/api/personas', methods=['POST'])
def create_persona():
    if not current_user.is_authenticated:
        return jsonify({"error": "Authentication required to create custom personas"}), 401
    try:
        data = request.get_json()
        name = data.get('name', '').strip()
        system_prompt = data.get('system_prompt', '').strip()
        icon = data.get('icon', 'sparkles').strip()
        if not name or not system_prompt:
            return jsonify({"error": "Name and system prompt are required"}), 400
        new_persona = Persona(
            name=name,
            system_prompt=system_prompt,
            icon=icon,
            is_system=False,
            user_id=current_user.id
        )
        db.session.add(new_persona)
        db.session.commit()
        return jsonify(new_persona.to_dict()), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/personas/<persona_id>', methods=['DELETE'])
@login_required
def delete_persona(persona_id):
    try:
        persona = Persona.query.filter_by(id=persona_id, user_id=current_user.id, is_system=False).first()
        if not persona:
            return jsonify({"error": "Custom persona not found or unauthorized"}), 404
        db.session.delete(persona)
        db.session.commit()
        return jsonify({"success": True, "message": "Custom persona deleted successfully"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# === RAG Helpers & Endpoints ===

def extract_text_from_file(file_bytes, filename):
    import os
    _, ext = os.path.splitext(filename.lower())
    text = ""
    try:
        if ext == '.pdf':
            from pypdf import PdfReader
            import io
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif ext in ['.docx', '.doc']:
            from docx import Document as DocxDocument
            import io
            docx_file = io.BytesIO(file_bytes)
            doc = DocxDocument(docx_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif ext == '.csv':
            import io
            import csv
            csv_file = io.StringIO(file_bytes.decode('utf-8', errors='ignore'))
            reader = csv.reader(csv_file)
            for row in reader:
                text += ", ".join(row) + "\n"
        else:
            text = file_bytes.decode('utf-8', errors='ignore')
    except Exception as e:
        print(f"Error parsing file {filename}: {e}")
    return text.strip()

def cosine_similarity(v1, v2):
    import math
    try:
        dot_product = sum(a * b for a, b in zip(v1, v2))
        magnitude1 = math.sqrt(sum(a * a for a in v1))
        magnitude2 = math.sqrt(sum(b * b for b in v2))
        if not magnitude1 or not magnitude2:
            return 0.0
        return dot_product / (magnitude1 * magnitude2)
    except Exception as e:
        return 0.0

def chunk_text(text, chunk_size=800, overlap=150):
    chunks = []
    if not text:
        return chunks
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks

def get_ollama_embedding(text, model_name=DEFAULT_MODEL):
    try:
        url = OLLAMA_API_URL.replace("/generate", "/embeddings")
        payload = {
            "model": model_name,
            "prompt": text
        }
        res = requests.post(url, json=payload, timeout=20)
        if res.ok:
            return res.json().get("embedding")
        
        # Fallback to newer /api/embed API
        embed_url = OLLAMA_API_URL.replace("/generate", "/embed")
        new_payload = {
            "model": model_name,
            "input": [text]
        }
        res = requests.post(embed_url, json=new_payload, timeout=20)
        if res.ok:
            embeddings = res.json().get("embeddings")
            if embeddings and len(embeddings) > 0:
                return embeddings[0]
    except Exception as e:
        print(f"Error calling Ollama embedding endpoint: {e}")
    return None

@app.route('/api/upload', methods=['POST'])
def upload_file():
    try:
        data = request.get_json()
        session_id = data.get('session_id')
        filename = data.get('filename')
        file_data = data.get('data')
        model_name = data.get('model', DEFAULT_MODEL).strip()
        
        if not session_id or not filename or not file_data:
            return jsonify({"error": "Missing session_id, filename, or data"}), 400
            
        user_id = current_user.id if current_user.is_authenticated else None
        if user_id:
            session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()
        else:
            session = ChatSession.query.filter_by(id=session_id, user_id=None).first()
            
        if not session:
            return jsonify({"error": "Session not found"}), 404
            
        import base64
        if ',' in file_data:
            file_data = file_data.split(',')[1]
        decoded = base64.b64decode(file_data)
        
        text = extract_text_from_file(decoded, filename)
        if not text:
            return jsonify({"error": "Could not extract text from document"}), 400
            
        import os
        _, ext = os.path.splitext(filename.lower())
        new_doc = Document(
            session_id=session_id,
            filename=filename,
            file_type=ext
        )
        db.session.add(new_doc)
        db.session.commit()
        
        chunks = chunk_text(text, chunk_size=800, overlap=150)
        
        chunk_objects = []
        for chunk_content in chunks:
            embedding = get_ollama_embedding(chunk_content, model_name)
            if not embedding:
                embedding = [0.0] * 512 # Fallback
                
            new_chunk = DocumentChunk(
                document_id=new_doc.id,
                session_id=session_id,
                content=chunk_content
            )
            new_chunk.embedding = embedding
            chunk_objects.append(new_chunk)
            
        db.session.bulk_save_objects(chunk_objects)
        db.session.commit()
        
        return jsonify({
            "success": True,
            "document": new_doc.to_dict(),
            "chunks_count": len(chunks)
        }), 201
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/sessions/<session_id>/documents', methods=['GET'])
def get_session_documents(session_id):
    user_id = current_user.id if current_user.is_authenticated else None
    if user_id:
        session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()
    else:
        session = ChatSession.query.filter_by(id=session_id, user_id=None).first()
        
    if not session:
        return jsonify({"error": "Session not found"}), 404
        
    docs = Document.query.filter_by(session_id=session_id).all()
    return jsonify({"documents": [d.to_dict() for d in docs]})

@app.route('/api/documents/<document_id>', methods=['DELETE'])
def delete_document(document_id):
    try:
        doc = Document.query.get(document_id)
        if not doc:
            return jsonify({"error": "Document not found"}), 404
        db.session.delete(doc)
        db.session.commit()
        return jsonify({"success": True, "message": "Document deleted successfully"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# === Export Endpoints ===

@app.route('/api/sessions/<session_id>/export/markdown', methods=['GET'])
def export_markdown(session_id):
    try:
        user_id = current_user.id if current_user.is_authenticated else None
        if user_id:
            session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()
        else:
            session = ChatSession.query.filter_by(id=session_id, user_id=None).first()
        
        if not session:
            return jsonify({"error": "Session not found"}), 404
        
        messages = ChatMessage.query.filter_by(session_id=session_id).order_by(ChatMessage.created_at.asc()).all()
        
        md_content = f"# {session.title or 'Chat Conversation'}\n\n"
        md_content += f"*Exported on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}*\n\n---\n\n"
        
        for msg in messages:
            if msg.role == 'user':
                md_content += f"## 🧑 User\n\n{msg.content}\n\n---\n\n"
            else:
                md_content += f"## 🤖 Assistant\n\n{msg.content}\n\n---\n\n"
        
        buffer = io.BytesIO()
        buffer.write(md_content.encode('utf-8'))
        buffer.seek(0)
        
        safe_title = "".join(c for c in (session.title or 'chat') if c.isalnum() or c in ' -_').strip()[:50]
        filename = f"{safe_title}.md"
        
        return send_file(buffer, mimetype='text/markdown', as_attachment=True, download_name=filename)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/sessions/<session_id>/export/pdf', methods=['GET'])
def export_pdf(session_id):
    try:
        user_id = current_user.id if current_user.is_authenticated else None
        if user_id:
            session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()
        else:
            session = ChatSession.query.filter_by(id=session_id, user_id=None).first()
        
        if not session:
            return jsonify({"error": "Session not found"}), 404
        
        messages = ChatMessage.query.filter_by(session_id=session_id).order_by(ChatMessage.created_at.asc()).all()
        
        from fpdf import FPDF
        import re
        
        class ChatPDF(FPDF):
            def header(self):
                self.set_font('Helvetica', 'B', 14)
                self.set_text_color(99, 102, 241)
                self.cell(0, 10, 'Privacy LLM - Chat Export', new_x='LMARGIN', new_y='NEXT', align='C')
                self.set_draw_color(226, 232, 240)
                self.line(10, self.get_y(), 200, self.get_y())
                self.ln(5)
            
            def footer(self):
                self.set_y(-15)
                self.set_font('Helvetica', 'I', 8)
                self.set_text_color(148, 163, 184)
                self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')
        
        pdf = ChatPDF()
        pdf.alias_nb_pages()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()
        
        # Title
        pdf.set_font('Helvetica', 'B', 16)
        pdf.set_text_color(15, 23, 42)
        title_text = session.title or 'Chat Conversation'
        pdf.cell(0, 12, title_text, new_x='LMARGIN', new_y='NEXT')
        
        pdf.set_font('Helvetica', '', 9)
        pdf.set_text_color(148, 163, 184)
        pdf.cell(0, 6, f"Exported on {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", new_x='LMARGIN', new_y='NEXT')
        pdf.ln(8)
        
        for msg in messages:
            is_user = msg.role == 'user'
            
            # Role label
            if is_user:
                pdf.set_fill_color(59, 130, 246)
                pdf.set_text_color(255, 255, 255)
                pdf.set_font('Helvetica', 'B', 10)
                pdf.cell(28, 7, '  USER', fill=True, new_x='LMARGIN', new_y='NEXT')
            else:
                pdf.set_fill_color(139, 92, 246)
                pdf.set_text_color(255, 255, 255)
                pdf.set_font('Helvetica', 'B', 10)
                pdf.cell(38, 7, '  ASSISTANT', fill=True, new_x='LMARGIN', new_y='NEXT')
            
            pdf.ln(3)
            
            # Clean markdown for PDF
            content = msg.content
            content = re.sub(r'```[\w]*\n?', '--- code ---\n', content)
            content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
            content = re.sub(r'\*(.+?)\*', r'\1', content)
            content = re.sub(r'`(.+?)`', r'\1', content)
            # Remove non-latin1 characters for fpdf compatibility
            content = content.encode('latin-1', 'replace').decode('latin-1')
            
            pdf.set_text_color(30, 41, 59)
            pdf.set_font('Helvetica', '', 10)
            pdf.multi_cell(0, 5.5, content)
            pdf.ln(4)
            
            # Divider line
            pdf.set_draw_color(226, 232, 240)
            y_pos = pdf.get_y()
            pdf.line(10, y_pos, 200, y_pos)
            pdf.ln(6)
        
        buffer = io.BytesIO()
        pdf.output(buffer)
        buffer.seek(0)
        
        safe_title = "".join(c for c in (session.title or 'chat') if c.isalnum() or c in ' -_').strip()[:50]
        filename = f"{safe_title}.pdf"
        
        return send_file(buffer, mimetype='application/pdf', as_attachment=True, download_name=filename)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# === Chat Routes ===

@app.route('/history', methods=['GET'])
def get_all_history():
    if not current_user.is_authenticated:
        # Guests don't get persistent history list in sidebar for now
        # Could implement cookie-based if needed, but per request implies simple guest chat
        return jsonify({"sessions": []})
        
    # Only show current user's sessions
    sessions = ChatSession.query.filter_by(user_id=current_user.id).order_by(ChatSession.created_at.desc()).all()
    session_list = []
    for s in sessions:
        session_list.append({
            "id": s.id,
            "title": s.title,
            "created_at": s.created_at.isoformat()
        })
    return jsonify({"sessions": session_list})

@app.route('/history/<session_id>', methods=['GET'])
def get_session_details(session_id):
    if current_user.is_authenticated:
        session = ChatSession.query.filter_by(id=session_id, user_id=current_user.id).first()
    else:
        # Guests can access sessions if they know ID? Only if it has NO user_id
        session = ChatSession.query.filter_by(id=session_id, user_id=None).first()
        
    if session:
        return jsonify(session.to_dict())
    return jsonify({"error": "Session not found"}), 404

@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        user_prompt = data.get('prompt', '').strip()
        model_name = data.get('model', DEFAULT_MODEL).strip()
        session_id = data.get('session_id')
        attachments = data.get('attachments', []) 
        is_image_gen = data.get('isImageGeneration', False)
        
        # Ingest new Persona & Slider options
        persona_id = data.get('persona_id')
        temperature = data.get('temperature')
        top_p = data.get('top_p')
        top_k = data.get('top_k')
        repeat_penalty = data.get('repeat_penalty')
        num_predict = data.get('num_predict')
        
        if not user_prompt and not attachments:
            return jsonify({"error": "Prompt or attachment required"}), 400

        # RAG-eligible document extensions
        rag_extensions = {'.pdf', '.docx', '.doc', '.csv', '.txt'}
        # Code/text file extensions (direct injection)
        code_extensions = {
            '.md', '.py', '.js', '.html', '.css', '.json', '.xml', 
            '.yaml', '.yml', '.log', '.sql', '.conf', '.ini', '.env',
            '.java', '.cpp', '.c', '.h', '.rb', '.go', '.rs', '.php', '.sh',
            '.bat', '.ps1', '.jsx', '.tsx', '.ts', '.vue', '.svelte'
        }
        
        # Process file attachments
        file_contents = []
        rag_embedded_files = []  # Track files that were RAG-embedded
        if attachments:
            for attachment in attachments:
                try:
                    filename = attachment.get('name', 'unknown')
                    file_data = attachment.get('data', '')
                    
                    _, ext = os.path.splitext(filename.lower())
                    
                    if ',' in file_data:
                        file_data_b64 = file_data.split(',')[1]
                    else:
                        file_data_b64 = file_data
                    
                    decoded_bytes = base64.b64decode(file_data_b64)
                    
                    if ext in rag_extensions:
                        # RAG pipeline: extract → chunk → embed → store
                        text = extract_text_from_file(decoded_bytes, filename)
                        if text:
                            rag_embedded_files.append({
                                'filename': filename,
                                'ext': ext,
                                'text': text
                            })
                            # Also add a summary to direct context for immediate response
                            summary = text[:2000] if len(text) > 2000 else text
                            file_contents.append({
                                'filename': filename,
                                'content': f"[Document uploaded and indexed for RAG]\nPreview: {summary}"
                            })
                    elif ext in code_extensions:
                        # Direct text injection for code/text files
                        decoded = decoded_bytes.decode('utf-8', errors='ignore')
                        
                        if len(decoded) > 102400:
                            decoded = decoded[:102400] + "\n... (truncated)"
                        
                        file_contents.append({
                            'filename': filename,
                            'content': decoded
                        })
                except Exception as e:
                    print(f"Error processing attachment {attachment.get('name', '?')}: {e}")
                    continue
        
        # Format file contents into prompt
        if file_contents:
            files_text = "\n\n"
            for fc in file_contents:
                files_text += f"--- File: {fc['filename']} ---\n{fc['content']}\n--- End of File ---\n\n"
            user_prompt = f"{files_text}User question: {user_prompt if user_prompt else 'Please analyze the attached files.'}"

        # Handle Image Gen (Stub)
        if is_image_gen:
             def generate_message_stub():
                msg = "**System:** Image Generation is not supported by the local Ollama backend presently."
                yield msg
             return Response(stream_with_context(generate_message_stub()), mimetype='text/plain')

        # Session Management
        session = None
        user_id = current_user.id if current_user.is_authenticated else None
        
        if session_id:
            if user_id:
                session = ChatSession.query.filter_by(id=session_id, user_id=user_id).first()
            else:
                session = ChatSession.query.filter_by(id=session_id, user_id=None).first()

        if not session:
            title = user_prompt[:40] + "..." if len(user_prompt) > 40 else user_prompt
            new_session_id = str(uuid.uuid4())
            session = ChatSession(id=new_session_id, user_id=user_id, title=title)
            db.session.add(session)
            db.session.commit()
            session_id = session.id
        
        # Process RAG-eligible attachments AFTER session exists
        if rag_embedded_files:
            for rag_file in rag_embedded_files:
                try:
                    new_doc = Document(
                        session_id=session_id,
                        filename=rag_file['filename'],
                        file_type=rag_file['ext']
                    )
                    db.session.add(new_doc)
                    db.session.commit()
                    
                    chunks = chunk_text(rag_file['text'], chunk_size=800, overlap=150)
                    chunk_objects = []
                    for chunk_content in chunks:
                        embedding = get_ollama_embedding(chunk_content, model_name)
                        if not embedding:
                            embedding = [0.0] * 512
                        new_chunk = DocumentChunk(
                            document_id=new_doc.id,
                            session_id=session_id,
                            content=chunk_content
                        )
                        new_chunk.embedding = embedding
                        chunk_objects.append(new_chunk)
                    
                    db.session.bulk_save_objects(chunk_objects)
                    db.session.commit()
                    print(f"RAG embedded: {rag_file['filename']} ({len(chunks)} chunks)")
                except Exception as e:
                    print(f"Error RAG-embedding {rag_file['filename']}: {e}")
            
        # Get selected system persona
        system_prompt = "You are a helpful, friendly, and knowledgeable AI assistant."
        if persona_id:
            persona = Persona.query.get(persona_id)
            if persona:
                system_prompt = persona.system_prompt

        # Ollama Payload
        ollama_payload = {
            "model": model_name,
            "prompt": user_prompt,
            "stream": True
        }
        
        # Map advanced option parameters to options payload
        options = {}
        if temperature is not None:
            options["temperature"] = float(temperature)
        if top_p is not None:
            options["top_p"] = float(top_p)
        if top_k is not None:
            options["top_k"] = int(top_k)
        if repeat_penalty is not None:
            options["repeat_penalty"] = float(repeat_penalty)
        if num_predict is not None:
            options["num_predict"] = int(num_predict)
            
        if options:
            ollama_payload["options"] = options
        
        # Check if RAG document chunks exist for this session
        rag_context = ""
        try:
            session_chunks = DocumentChunk.query.filter_by(session_id=session_id).all()
            if session_chunks and user_prompt:
                prompt_embedding = get_ollama_embedding(user_prompt, model_name)
                if prompt_embedding:
                    chunk_scores = []
                    for chunk in session_chunks:
                        score = cosine_similarity(prompt_embedding, chunk.embedding)
                        chunk_scores.append((score, chunk.content))
                    
                    chunk_scores.sort(key=lambda x: x[0], reverse=True)
                    
                    # Retrieve top 3 chunks with positive score
                    top_chunks = [c for s, c in chunk_scores if s > 0.1][:3]
                    if top_chunks:
                        rag_context += "\n[CONTEXT FROM UPLOADED DOCUMENTS]\n"
                        for content in top_chunks:
                            rag_context += f"- {content}\n"
                        rag_context += "[END OF CONTEXT]\n\n"
        except Exception as e:
            print(f"Error performing RAG vector search: {e}")

        # Build Context from DB
        recent_msgs = ChatMessage.query.filter_by(session_id=session_id).order_by(ChatMessage.created_at.desc()).limit(5).all()
        recent_msgs.reverse() 
        
        # Compile prompts prepended with System Persona instructions
        full_context_prompt = f"System: {system_prompt}\n\n"
        for msg in recent_msgs:
            role_display = "User" if msg.role == "user" else "Assistant"
            full_context_prompt += f"{role_display}: {msg.content}\n"
            
        # Enrich the final user prompt sent to LLM with RAG context
        enriched_user_prompt = user_prompt
        if rag_context:
            enriched_user_prompt = f"{rag_context}User question: {user_prompt}"
            
        full_context_prompt += f"User: {enriched_user_prompt}\nAssistant:"
        
        ollama_payload['prompt'] = full_context_prompt

        def generate():
            full_response = ""
            prompt_tokens = 0
            eval_tokens = 0
            
            try:
                with requests.post(OLLAMA_API_URL, json=ollama_payload, stream=True, timeout=120) as response:
                    response.raise_for_status()
                    
                    for line in response.iter_lines():
                        if line:
                            try:
                                json_response = json.loads(line)
                                if 'response' in json_response:
                                    chunk = json_response['response']
                                    full_response += chunk
                                    yield chunk
                                
                                # Extract token usage on completion chunk
                                if json_response.get('done') is True:
                                    prompt_tokens = json_response.get('prompt_eval_count', 0)
                                    eval_tokens = json_response.get('eval_count', 0)
                            except:
                                pass
                
                # Yield the token usage information at the very end of the stream
                if prompt_tokens > 0 or eval_tokens > 0:
                    yield f"\n\n[TOKEN_USAGE:prompt_tokens={prompt_tokens},eval_tokens={eval_tokens}]"
                
                # Update DB with full conversation pair and token counts
                with app.app_context():
                     # Save user message
                     user_msg = ChatMessage(session_id=session_id, role="user", content=user_prompt)
                     db.session.add(user_msg)
                     # Save assistant message
                     asst_msg = ChatMessage(
                         session_id=session_id, 
                         role="assistant", 
                         content=full_response, 
                         prompt_tokens=prompt_tokens, 
                         eval_tokens=eval_tokens
                     )
                     db.session.add(asst_msg)
                     db.session.commit()
                    
            except Exception as e:
                yield f"\n[Error: {str(e)}]"
        
        return Response(stream_with_context(generate()), mimetype='text/plain', headers={"X-Session-ID": session_id})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

def seed_personas():
    try:
        default_presets = [
            Persona(
                name="General Assistant",
                system_prompt="You are a helpful, friendly, and highly intelligent general assistant. You provide clear, accurate, and context-aware responses with excellent structure. Balance completeness and conciseness, use clear formatting (bolding, lists) to emphasize key details, and adapt your tone to be professional yet warm.",
                icon="sparkles",
                is_system=True
            ),
            Persona(
                name="Expert Software Engineer",
                system_prompt="You are an elite, highly experienced Senior Software Engineer and Architect. You provide exceptionally clean, secure, performant, and well-commented code following the industry's best practices (SOLID principles, DRY, modular design). Always explain your architectural and implementation decisions, point out potential edge cases, write comprehensive docstrings and comments, and present code inside nicely formatted markdown code blocks.",
                icon="code",
                is_system=True
            ),
            Persona(
                name="Scientific Researcher",
                system_prompt="You are a distinguished scientific researcher and academic writer. You analyze queries with rigorous critical thinking, scientific objectivity, and academic precision. Formulate balanced, data-driven, and objective explanations. Reference established theoretical frameworks, research methods, and mathematical expressions where appropriate. Maintain a formal, authoritative, and objective scientific tone.",
                icon="microscope",
                is_system=True
            ),
            Persona(
                name="Copywriter / Editor",
                system_prompt="You are an expert copywriter, creative content strategist, and meticulously detail-oriented editor. You refine text to make it exceptionally engaging, persuasive, clear, and grammatically flawless, maintaining the author's original voice while elevating the impact. Suggest structural and stylistic enhancements, and provide alternative headline or phrasing choices to maximize audience engagement.",
                icon="pen-tool",
                is_system=True
            ),
            Persona(
                name="Concise Summarizer",
                system_prompt="You are an expert executive summarizer and data synthesist. You distill complex, dense inputs into their absolute core ideas, key arguments, logical takeaways, and actionable bullet-point steps. Structure your summaries with clear markdown headers, bold vital terms, eliminate all conversational filler, fluff, or preambles, and present information in a high-density, scannable format.",
                icon="align-left",
                is_system=True
            )
        ]
        
        # Upgrades seeded system personas prompts dynamically on start if already exist
        for preset in default_presets:
            existing = Persona.query.filter_by(name=preset.name, is_system=True).first()
            if existing:
                existing.system_prompt = preset.system_prompt
                existing.icon = preset.icon
            else:
                db.session.add(preset)
        
        db.session.commit()
        print("Successfully seeded/updated system personas!")
    except Exception as e:
        print(f"Error seeding personas: {e}")

if __name__ == '__main__':
    with app.app_context():
        # Automated Schema Migration for local SQLite DB schema upgrades
        try:
            db.session.execute(db.text("ALTER TABLE chat_message ADD COLUMN prompt_tokens INTEGER"))
            db.session.execute(db.text("ALTER TABLE chat_message ADD COLUMN eval_tokens INTEGER"))
            db.session.commit()
            print("Successfully migrated local SQLite DB schema for token usage tracking!")
        except Exception:
            db.session.rollback() # Columns already exist, skip migration
            
        db.create_all() # Re-create tables if DB file deleted
        seed_personas()
    app.run(host='0.0.0.0', port=5000, debug=True)

